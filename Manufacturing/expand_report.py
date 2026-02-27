"""
Expand the Drilling and Milling sections of Lab 1 report.
Replaces paragraphs P93 onwards (section 2) with fully expanded content,
then re-appends the Conclusions section.
"""
from docx import Document
from docx.shared import Pt
from copy import deepcopy
import shutil
import os

SRC = r"C:\Users\stani\OneDrive\Pulpit\PŁ_dokumenty\Manufacturing\Reports\Report L1 Survey of machine tools for selected cutting processes.docx"
BACKUP = SRC.replace(".docx", " - BACKUP.docx")

# Create backup
shutil.copy2(SRC, BACKUP)
print(f"Backup created: {BACKUP}")

doc = Document(SRC)

# We will delete paragraphs from P92 (empty heading) to P147 (end)
# and replace them with our expanded content.
# Keep paragraphs P0..P91 (the entire Lathe section).

# Count total paragraphs
total = len(doc.paragraphs)
print(f"Total paragraphs before edit: {total}")

# Strategy: delete paragraphs from index 92 to end (backwards), 
# then add new content.

# Delete from end backwards
for i in range(total - 1, 91, -1):
    p = doc.paragraphs[i]
    p_element = p._element
    p_element.getparent().remove(p_element)

print(f"Paragraphs after deletion: {len(doc.paragraphs)}")

# Helper function to add a paragraph with style
def add_para(text, style='Normal'):
    p = doc.add_paragraph(text, style=style)
    return p

def add_heading(text, level=1):
    return doc.add_heading(text, level=level)

# ============================================================
# SECTION 2: DRILLING AND MILLING MACHINES
# ============================================================

add_heading('2. Drilling and Milling Machines', level=1)
add_para('')

# ============================================================
# A. DRILLING MACHINES
# ============================================================
add_heading('A. Drilling Machines', level=2)

# --- Structure and Construction ---
add_heading('Structure and Construction', level=3)
add_para(
    'The bench or column drilling machine is one of the most fundamental machine tools found in any workshop. '
    'Its construction is organized around a robust vertical column, typically made of cast iron, which connects '
    'the heavy base at the bottom to the headstock at the top. The base serves as a stable foundation and often '
    'features T-slots or bolt holes for securing large workpieces or fixtures directly. The headstock houses the '
    'electric motor, the V-belt or gear drive transmission, and the spindle mechanism. The spindle is a hollow, '
    'rotating shaft that receives the cutting tool, and it can be advanced vertically downward toward the workpiece '
    'by means of a hand-operated lever (quill feed) or, on larger machines, an automatic power feed mechanism.'
)
add_para(
    'Between the base and the headstock, the work table is mounted on the column and can be adjusted vertically '
    'to accommodate workpieces of varying height. The table surface is typically flat and machined with T-slots '
    'for clamping vices, jigs, or fixtures. On some models, the table can also be tilted for angular drilling '
    'operations. Depth stops are provided on the quill to control the depth of the drilled hole precisely. '
    'A key characteristic of the standard column drill press is that the workpiece must be positioned manually '
    'beneath the spindle, since neither the table nor the spindle can be moved horizontally.'
)

add_para('Figure 7. A sketch of a drilling machine (bench drill press).', style='Caption')
add_para('Legend:')
add_para('1. Column')
add_para('2. Base')
add_para('3. Work Table (with T-slots)')
add_para('4. Headstock (with motor and drive)')
add_para('5. Spindle and Quill')
add_para('6. Hand Feed Lever')
add_para('7. Depth Stop')
add_para('8. Drill Chuck with Drill Bit')
add_para('')

# --- Kinematics ---
add_heading('Kinematics and Movements', level=3)
add_para(
    'The kinematic system of a drilling machine is fundamentally different from that of a lathe. '
    'In a drilling machine, the cutting tool performs both of the primary movements. The drill bit executes '
    'the main rotary cutting motion (rotation around its own axis), and simultaneously performs the linear '
    'feed motion (axial advancement into the material). The workpiece, in contrast, remains completely '
    'stationary throughout the entire process, being clamped rigidly to the table or held in a vice.'
)
add_para(
    'This kinematic arrangement — where the tool provides both rotation and feed — differs significantly '
    'from the lathe, where the workpiece rotates and the tool feeds, and from the milling machine, where '
    'the tool rotates but the workpiece provides the feed motion. The rotational speed of the spindle '
    '(measured in revolutions per minute, RPM) and the feed rate (measured in millimeters per revolution, mm/rev) '
    'are the two primary cutting parameters. The spindle speed is selected based on the cutting speed formula: '
    'n = (1000 × Vc) / (π × d), where Vc is the recommended cutting speed for the material and d is the drill '
    'diameter. The feed rate is chosen depending on the drill size, material hardness, and required hole quality. '
    'Excessive feed can lead to drill breakage, while too low a feed causes rubbing and work hardening rather '
    'than effective cutting.'
)
add_para(
    'On the sketch below, the arrows indicate the directions of movement: a circular arrow around the spindle '
    'axis for the main cutting motion, and a straight downward arrow for the feed motion. The workpiece '
    'is shown fixed to the table with no motion arrows.'
)
add_para('')

# --- Surfaces and Involved Elements ---
add_heading('Shaped Surfaces and Involved Elements', level=3)
add_para(
    'Drilling machines are primarily designed to produce cylindrical holes of various diameters and depths. '
    'However, the range of machinable surfaces extends well beyond simple through-holes. Drilling operations '
    'create blind holes (holes that do not pass through the entire workpiece), through-holes, and stepped holes '
    'of varying diameters. Beyond basic drilling, several secondary operations can be performed on the same machine:'
)
add_para(
    '• Reaming — a finishing operation that uses a multi-fluted reamer to improve the dimensional accuracy, '
    'roundness, and surface finish of a previously drilled hole. Reaming typically removes only 0.1–0.3 mm '
    'of material and produces tolerances of IT7 or better.'
)
add_para(
    '• Counterboring — creating a flat-bottomed cylindrical recess at the top of a hole to accommodate the '
    'head of a socket-head cap screw or bolt, allowing it to sit flush with or below the workpiece surface.'
)
add_para(
    '• Countersinking — producing a conical recess (typically at 60°, 82°, or 90°) at the entrance of a '
    'hole to accommodate flat-head screws or to deburr the hole edge.'
)
add_para(
    '• Tapping — cutting internal threads inside a drilled hole using a tap. This can be done by hand using '
    'a tap wrench or by using a tapping attachment on the drilling machine that reverses the spindle at the '
    'bottom of the stroke.'
)
add_para(
    '• Spot facing — machining a small, flat, circular surface around a hole on a rough or curved surface to '
    'provide a proper seating area for a washer or bolt head.'
)
add_para(
    'In all these operations, the spindle and the tool it carries perform the primary rotary motion. The '
    'vertical feed of the quill advances the tool into the workpiece. The table and the fixturing clamp the '
    'workpiece securely to resist the cutting forces, particularly the torque and the axial thrust force generated '
    'during drilling.'
)
add_para('')

# --- Fixing and Clamping Methods ---
add_heading('Fixing and Clamping Methods', level=3)
add_para(
    'Proper clamping of the workpiece is critical in drilling operations to resist the torque and thrust forces '
    'produced by the drill. Unlike a lathe, where the workpiece rotates, in drilling the workpiece must be held '
    'absolutely still. Several methods are commonly used:'
)
add_para(
    '• Machine Vice — the most common clamping device for small to medium-sized workpieces. The vice is bolted '
    'to the table through the T-slots. The workpiece is placed between the fixed and movable jaws and tightened '
    'securely. V-blocks can be placed inside the vice jaws to hold cylindrical workpieces. For precision work, '
    'the vice can be aligned with the spindle axis using a dial indicator.'
)
add_para(
    '• T-slot Clamps and Bolts — for larger or irregularly shaped workpieces that cannot fit in a vice, strap '
    'clamps, step blocks, and T-bolts are used to secure the part directly to the table. This method offers '
    'the most flexibility in terms of workpiece geometry and provides very rigid support.'
)
add_para(
    '• Drilling Jigs — specialized fixtures designed for batch or mass production. A drilling jig positions and '
    'clamps the workpiece and incorporates hardened drill bushings that guide the drill bit to the exact required '
    'location. This eliminates the need for marking out and ensures consistent hole placement across many parts, '
    'significantly reducing setup time and improving accuracy.'
)
add_para(
    '• Hand Holding — for very quick, non-critical work on small pieces, the workpiece may sometimes be held by '
    'hand. However, this practice is generally discouraged as the drill can grab the workpiece and spin it, '
    'presenting a serious safety hazard. A piece of scrap wood is often placed beneath the workpiece to protect '
    'the table and provide support at the hole exit.'
)
add_para(
    '• Angle Plates — when holes must be drilled at specific angles relative to a datum surface, angle plates '
    'or adjustable angle vices are used to orient the workpiece at the required inclination under the vertical spindle.'
)

add_para('Figure 9. A sketch of Fixing Methods on a Drilling Machine.', style='Caption')
add_para('Legend:')
add_para('1. Machine vice bolted to table')
add_para('2. T-slot clamps with step blocks')
add_para('3. Drilling jig with hardened bushings')
add_para('4. V-block for cylindrical parts')
add_para('')

# --- Tools and Mounting ---
add_heading('Tools and Mounting', level=3)
add_para(
    'The cutting tools used on drilling machines are multi-point rotary tools, as opposed to the single-point '
    'tools used on a lathe. The most common tool is the twist drill, a cylindrical tool with two helical flutes '
    'that create the cutting edges at the tip and provide channels for chip evacuation. Twist drills are available '
    'with cylindrical shanks (for smaller diameters, held in a drill chuck) or tapered shanks (for larger '
    'diameters, inserted directly into the spindle taper, typically a Morse taper).'
)
add_para(
    'Other important tools include:'
)
add_para(
    '• Center Drills — short, rigid drills used to create a starting point (center hole) for subsequent drilling '
    'to prevent the twist drill from wandering off position.'
)
add_para(
    '• Reamers — multi-fluted cylindrical or tapered tools for finishing holes to tight tolerances. '
    'They have straight or helical flutes and are designed for very light material removal.'
)
add_para(
    '• Counterbore Cutters — tools with a pilot that guides them into an existing hole, while the larger '
    'cutting diameter machines the flat-bottomed recess.'
)
add_para(
    '• Countersink Bits — conical tools used to create chamfered entries to holes.'
)
add_para(
    '• Taps — used with a tapping head or reversible spindle to cut internal threads.'
)
add_para(
    'All cylindrical-shank tools are mounted in a keyed or keyless drill chuck, which is itself held in the '
    'spindle by a Morse taper arbor. Taper-shank tools are inserted directly into the spindle bore. Removal of '
    'taper-shank tools is accomplished using a drift key inserted through a slot in the spindle. Quick-change '
    'toolholders are available for production environments to minimize tool changeover time.'
)

add_para('Figure 10. A sketch of Tools and Mounting on a Drilling Machine.', style='Caption')
add_para('Legend:')
add_para('1. Twist drill (cylindrical shank) in drill chuck')
add_para('2. Twist drill (taper shank) in spindle')
add_para('3. Center drill')
add_para('4. Reamer')
add_para('5. Counterbore with pilot')
add_para('6. Countersink')
add_para('7. Tap (for threading)')
add_para('')

# --- Special Equipment and Types ---
add_heading('Special Equipment and Drilling Machine Types', level=3)
add_para(
    'Standard column drilling machines can be enhanced with various accessories to increase their capabilities. '
    'A cross-slide table (compound table) can be mounted on the drill press table to allow the workpiece to be '
    'moved precisely in the X and Y directions, transforming the drill press into a rudimentary milling or '
    'coordinate drilling machine. Tapping attachments with automatic spindle reversal are used for efficient '
    'thread-cutting operations. Coolant systems deliver cutting fluid to the drill point to reduce heat, '
    'extend tool life, and improve surface finish.'
)
add_para(
    'Several specialized types of drilling machines exist for different industrial applications:'
)
add_para(
    '• Radial Drilling Machine — designed for large, heavy workpieces that are difficult to move. The machine '
    'features a radial arm mounted on a vertical column. The arm can swing horizontally around the column, '
    'and the spindle head can be moved along the arm. This allows the drill to be positioned over any point on '
    'the workpiece without moving the part, making it ideal for machining large castings, frames, and structural '
    'members. The workpiece is placed on the base or the floor, and the arm is brought to the work.'
)
add_para(
    '• Gang Drilling Machine — consists of multiple drilling spindles mounted in a row on a single table. '
    'Each spindle can hold a different tool (e.g., center drill, twist drill, reamer), and the workpiece is '
    'moved sequentially from one spindle to the next. This arrangement is highly efficient for production work '
    'requiring multiple operations on the same hole.'
)
add_para(
    '• Multi-spindle Drilling Machine — features a special multi-spindle head with multiple drill spindles '
    'driven from a single power source through a system of gears. All spindles operate simultaneously, drilling '
    'many holes at once in a single stroke. The positions of the spindles can be adjusted to match the hole '
    'pattern of the workpiece. This type of machine is essential in mass production where cycle time must be minimized.'
)
add_para(
    '• Deep-Hole Drilling Machine — a specialized machine designed for drilling holes with a very high '
    'length-to-diameter ratio (typically greater than 10:1). These machines use special gun drills or BTA '
    '(Boring and Trepanning Association) drills with internal coolant supply to evacuate chips from deep holes. '
    'Applications include the drilling of gun barrels, hydraulic cylinders, and oil exploration equipment.'
)
add_para(
    '• CNC Drilling Machine — a computer numerically controlled machine that automates the positioning and '
    'drilling process. The table or spindle is moved automatically to programmed coordinates, enabling high '
    'precision and repeatability without manual intervention. CNC drilling machines are widely used in PCB '
    '(printed circuit board) manufacturing and aerospace applications.'
)
add_para('')
add_para('')

# ============================================================
# B. MILLING MACHINES
# ============================================================
add_heading('B. Milling Machines', level=2)

# --- Structure and Construction ---
add_heading('Structure and Construction', level=3)
add_para(
    'Milling machines are robust, versatile machine tools designed for heavy material removal and the generation '
    'of complex shapes. The most common configuration is the knee-and-column type, which is built around a '
    'massive column made of cast iron that houses the spindle drive train (motor, gearbox, and spindle). The '
    'column is the main structural backbone of the machine and provides rigidity against cutting forces.'
)
add_para(
    'The vertical knee slides along the front face of the column on precision guideways and can be raised or '
    'lowered to adjust the vertical position of the workpiece relative to the cutter. On top of the knee sits '
    'the saddle, which provides transverse (in-out) motion, and on top of the saddle is the work table, which '
    'provides longitudinal (left-right) motion. Together, these three degrees of freedom (vertical, transverse, '
    'and longitudinal) allow the workpiece to be positioned precisely relative to the rotating cutter.'
)
add_para(
    'Two principal orientations of milling machines exist. In a horizontal milling machine, the spindle axis '
    'is oriented horizontally, and an overarm extends from the top of the column to support the outer end of '
    'the arbor (a long shaft that holds the cutter). The arbor is supported by one or more arbor supports '
    '(yoke or bearing brackets) hanging from the overarm, which resist the cutting forces and prevent arbor '
    'deflection. In a vertical milling machine, the spindle axis is oriented perpendicular (vertical) to the '
    'table surface, and the cutter is mounted directly on the spindle nose using collets, chucks, or shell mill '
    'adapters. Some vertical milling machines have a swiveling head that allows the spindle to be tilted at an '
    'angle for angular milling operations.'
)

add_para('Figure 8. A sketch of a horizontal milling machine.', style='Caption')
add_para('Legend:')
add_para('1. Column (housing motor and gearbox)')
add_para('2. Overarm')
add_para('3. Arbor with Cutter and Arbor Support')
add_para('4. Work Table (with T-slots)')
add_para('5. Saddle')
add_para('6. Knee')
add_para('7. Base')
add_para('')

# --- Kinematics ---
add_heading('Kinematics and Movements', level=3)
add_para(
    'In milling, the main cutting movement is the rotation of the cutting tool, driven by the spindle. '
    'The feed movement is performed by the workpiece, which is clamped to the table and moved linearly '
    'relative to the rotating cutter. This feed can be in the X direction (longitudinal, along the table), '
    'the Y direction (transverse, across the saddle), or the Z direction (vertical, by raising or lowering '
    'the knee). This clear separation of rotation (tool) and translation (workpiece) is what distinguishes '
    'milling from drilling, where the tool performs both motions.'
)
add_para(
    'There are two fundamental modes of milling based on the relationship between the direction of cutter '
    'rotation and the direction of workpiece feed:'
)
add_para(
    '• Up-milling (conventional milling) — the cutter rotation opposes the feed direction. The chip starts '
    'thin and grows thicker. This mode tends to lift the workpiece from the table, requiring robust clamping, '
    'but it is gentler on the machine and is preferred for machines with backlash in the feed screws.'
)
add_para(
    '• Down-milling (climb milling) — the cutter rotation is in the same direction as the feed. The chip '
    'starts thick and becomes thinner. This provides a better surface finish, generates less heat at the cutting '
    'edge, and the cutting force pushes the workpiece down into the table. However, it requires a machine with '
    'backlash-free feed screws (ball screws) to prevent the cutter from pulling the workpiece into it.'
)
add_para(
    'The cutting speed (Vc) in milling is determined by the cutter diameter and spindle RPM: Vc = (π × d × n) / 1000. '
    'The feed rate is typically expressed as feed per tooth (fz), which is the amount of material each cutting edge '
    'removes per revolution. The table feed rate (Vf) equals: Vf = fz × z × n, where z is the number of teeth on '
    'the cutter and n is the RPM.'
)
add_para('')

# --- Surfaces and Involved Elements ---
add_heading('Shaped Surfaces and Involved Elements', level=3)
add_para(
    'Milling machines are among the most versatile machine tools, capable of producing a wide variety of surface '
    'geometries. The range of machinable surfaces includes:'
)
add_para(
    '• Flat horizontal surfaces (planes) — produced by face milling (using a face mill on a vertical machine) '
    'or by slab milling (using a cylindrical cutter on a horizontal machine). This is the most basic milling operation.'
)
add_para(
    '• Flat vertical surfaces (shoulders) — created by side milling or end milling. The tool cuts on its periphery '
    'and/or on its end face to produce a step or shoulder on the workpiece.'
)
add_para(
    '• Slots and grooves — produced using slot drills, end mills, or side-and-face cutters. T-slots (for machine '
    'tables), dovetail slots (for slideway connections), and keyways (for shaft-to-hub connections) are common '
    'examples of slotting operations.'
)
add_para(
    '• Angular surfaces and chamfers — created using single-angle or double-angle cutters, or by tilting the '
    'spindle head on a vertical milling machine.'
)
add_para(
    '• Complex 2D and 3D contours — achieved on machines with multiple axes of controlled motion. By coordinating '
    'the simultaneous movement of the table in the X and Y directions (and Z for 3D), complex profiles such as '
    'cam surfaces, die cavities, and mold shapes can be machined. This capability is greatly enhanced by CNC '
    '(computer numerical control).'
)
add_para(
    '• Gear teeth — spur, bevel, and helical gear teeth can be generated on a milling machine using specially '
    'shaped involute gear cutters and a dividing head (indexing head) to rotate the workpiece by the correct angle '
    'between each pass.'
)
add_para(
    'In all of these operations, the spindle rotates the cutter, and the table, saddle, and knee provide the '
    'three-axis positioning of the workpiece. The choice of cutter type, feed direction, and axis of motion '
    'determines the surface geometry that is created.'
)
add_para('')

# --- Fixing and Clamping Methods ---
add_heading('Fixing and Clamping Methods', level=3)
add_para(
    'Because milling generates significant cutting forces — including tangential, radial, and axial components — '
    'the workpiece must be clamped very securely to the table. Several common clamping methods are used on milling machines:'
)
add_para(
    '• Machine Vice — the standard clamping method for small to medium-sized workpieces. Milling vices are heavier '
    'and more rigid than drilling vices, with wide jaws and a solid construction designed to resist the high lateral '
    'cutting forces. The vice is bolted to the table through T-slots and can be aligned precisely using a dial '
    'indicator against the fixed jaw. Swivel-base vices allow the workpiece to be rotated to a specific angle '
    'for angular milling.'
)
add_para(
    '• T-slot Clamps, Bolts, and Step Blocks — for large, irregular, or awkward workpieces that do not fit in a '
    'vice, clamping kits with strap clamps, step blocks, jack screws, and T-bolts are used to fasten the part '
    'directly to the table. This method provides flexible and rigid clamping for virtually any shape of workpiece.'
)
add_para(
    '• Dividing Head (Indexing Head) — a precision device that holds and positions the workpiece for operations '
    'requiring accurate angular rotation, such as milling gear teeth, hexagonal forms, splines, or equally spaced '
    'slots around a circumference. The dividing head typically uses a worm gear mechanism with a 40:1 ratio, '
    'meaning 40 turns of the handle rotate the workpiece by one full revolution. Index plates with multiple hole '
    'circles allow for precise fractional turns.'
)
add_para(
    '• Rotary Table — a circular table that can be rotated continuously or indexed to specific angles. It is '
    'used for machining circular features such as arcs, curves, and bolt circles. When combined with CNC control, '
    'the rotary table effectively adds a fourth axis to the machine.'
)
add_para(
    '• Fixtures — special-purpose clamping devices designed for specific parts in batch or mass production. '
    'Fixtures locate the workpiece accurately and clamp it quickly, reducing setup time and ensuring repeatability '
    'across many parts. They are typically bolted to the table and may incorporate toggle clamps, hydraulic clamps, '
    'or pneumatic actuators for rapid operation.'
)
add_para(
    '• Magnetic Chucks — used for clamping flat ferromagnetic workpieces. An electromagnetic or permanent magnetic '
    'chuck holds the part by magnetic attraction, eliminating the need for mechanical clamps. This is particularly '
    'useful for surface grinding operations but can also be used in light milling.'
)

add_para('Figure 11. A sketch of Fixing Methods on a Milling Machine.', style='Caption')
add_para('Legend:')
add_para('1. Machine vice bolted to table via T-slots')
add_para('2. T-slot clamps with step blocks')
add_para('3. Dividing head with workpiece')
add_para('4. Rotary table')
add_para('')

# --- Tools and Mounting ---
add_heading('Tools and Mounting', level=3)
add_para(
    'Milling machines use multi-point rotary cutting tools, as opposed to the single-point tools used on lathes. '
    'The wide variety of cutters available is one of the reasons for the milling machine\'s versatility. The most '
    'common types of milling cutters include:'
)
add_para(
    '• Slab (Cylindrical) Mill — a plain cylindrical cutter with helical teeth on its periphery only. It is '
    'mounted on a horizontal arbor and is used for machining wide, flat surfaces (slab milling or peripheral milling). '
    'The helical teeth ensure a gradual engagement with the workpiece, reducing shock and vibration.'
)
add_para(
    '• Side-and-Face Cutter — a disc-shaped cutter with teeth on both the periphery and the side faces. It is used '
    'for milling slots, shoulders, and steps. When two or more are ganged together on an arbor, they can machine '
    'multiple surfaces simultaneously (straddle milling or gang milling).'
)
add_para(
    '• End Mill — a cylindrical cutter with cutting edges on both the end face and the periphery. End mills are the '
    'workhorses of vertical milling. They can perform plunge cutting (like a drill), peripheral cutting (along a '
    'contour), and face cutting. Available in a wide range of diameters, number of flutes (2, 3, 4, or more), and '
    'materials (HSS, solid carbide, cobalt). Two-flute end mills are designed for plunging and slotting; four-flute '
    'end mills provide a better surface finish for side and face milling.'
)
add_para(
    '• Face Mill — a large-diameter cutter with multiple replaceable carbide inserts mounted on its face. Face mills '
    'are used on vertical milling machines for rapid material removal and finishing of large flat surfaces. The inserts '
    'can be indexed (rotated) or replaced without removing the cutter body from the spindle.'
)
add_para(
    '• Slot Drill — similar to an end mill but with only two flutes and center-cutting capability. Designed '
    'specifically for plunge cutting and slot machining. The center-cutting design allows it to drill straight down '
    'like a drill bit before milling laterally.'
)
add_para(
    '• Form Cutters — specially shaped cutters for producing specific profiles, such as concave, convex, gear '
    'involute, corner-rounding, or thread-milling cutters. The shape of the cutter is the inverse of the desired '
    'profile on the workpiece.'
)
add_para(
    '• Fly Cutter — a single-point tool mounted in a rotating body. Despite having only one cutting edge, it can '
    'produce excellent surface finishes on large flat areas at low cost.'
)
add_para(
    'Tool mounting depends on the machine type. On a horizontal milling machine, cutters are mounted on a long '
    'arbor that passes through the spindle and is supported by the overarm. Spacer collars are used to position '
    'the cutter at the correct location along the arbor. On a vertical milling machine, tools are held in the '
    'spindle using collets (R8, ER, or other systems), drill chucks, shell mill adapters, or directly via a '
    'Morse taper. Quick-change toolholding systems are common in CNC milling machines to minimize tool changeover time.'
)

add_para('Figure 12. A sketch of Tools and Mounting on a Milling Machine.', style='Caption')
add_para('Legend:')
add_para('1. Slab mill on horizontal arbor')
add_para('2. Side-and-face cutter')
add_para('3. End mill in collet on vertical spindle')
add_para('4. Face mill with carbide inserts')
add_para('5. Slot drill')
add_para('6. Form cutter (gear tooth profile)')
add_para('')

# --- Special Equipment and Types ---
add_heading('Special Equipment and Milling Machine Types', level=3)
add_para(
    'Standard knee-and-column milling machines can be equipped with a range of accessories to extend their '
    'capabilities. Common special equipment includes:'
)
add_para(
    '• Dividing Head (Indexing Head) — as described in the clamping section, this is used for equally spacing '
    'cuts around a circumference, such as when milling gear teeth, splines, or hexagonal forms. It is one of the '
    'most important accessories for a milling machine.'
)
add_para(
    '• Slotting Attachment — converts the rotary output of the spindle into a reciprocating linear motion, '
    'allowing the milling machine to perform slotting operations similar to a slotter or shaper. This is useful '
    'for machining internal keyways and splines.'
)
add_para(
    '• Vertical Milling Attachment — mounted on a horizontal milling machine to convert it for vertical milling '
    'operations, adding versatility without requiring a separate vertical milling machine.'
)
add_para(
    '• Rack Milling Attachment — a specialized accessory for milling gear racks (straight teeth) that converts '
    'rotary motion into the linear indexing needed between successive tooth cuts.'
)
add_para(
    'In terms of specialized milling machine types, several important variants exist:'
)
add_para(
    '• Universal Milling Machine — essentially a horizontal milling machine with the addition of a table swivel '
    'that allows the table to be rotated in the horizontal plane. This is essential for helical milling operations, '
    'such as cutting helical gears, twist drills, and milling cutters themselves. A dividing head connected to '
    'the table lead screw via change gears provides the coordinated helical feed.'
)
add_para(
    '• Bed-Type Milling Machine — differs from the knee-and-column type in that the table is mounted on a fixed '
    'bed and can only move longitudinally (X). The spindle head moves vertically (Z) on the column. This design '
    'provides greater rigidity for heavy-duty milling of large castings and structural parts.'
)
add_para(
    '• Planer-Type Milling Machine (Plano-Miller) — a very large machine with multiple spindle heads mounted on '
    'cross-rails above a long table. Designed for machining very large workpieces such as machine tool beds, '
    'bridge components, and heavy frames. Multiple surfaces can be machined simultaneously using the different '
    'spindle heads.'
)
add_para(
    '• CNC Milling Machine / Machining Center — a computer-controlled milling machine that automates all axis '
    'movements, spindle speeds, and feed rates. Modern CNC machining centers feature automatic tool changers (ATC) '
    'with tool magazines holding 20–100+ tools, allowing multiple operations (drilling, tapping, milling, boring) '
    'to be performed in a single setup without manual intervention. Five-axis CNC machines add two rotary axes, '
    'enabling the machining of extremely complex 3D geometries such as turbine blades and aerospace components.'
)
add_para(
    '• Copy (Tracer) Milling Machine — uses a stylus that follows a template or model, and the cutter replicates '
    'the shape on the workpiece. This was an important method for producing complex 3D shapes before CNC became '
    'widely available. It is still used in some die and mold making applications.'
)
add_para('')
add_para('')

# ============================================================
# CONCLUSIONS
# ============================================================
add_heading('Conclusions', level=1)
add_para(
    'The laboratory survey highlighted the fundamental differences in kinematics between the three primary '
    'classes of machine tools studied. The lathe is unique in that the workpiece provides the main rotary motion '
    'while the tool provides the feed, making it the superior choice for creating radially symmetric parts '
    '(surfaces of revolution). In contrast, both drilling and milling machines rely on the tool for the main '
    'rotary cutting motion, but differ fundamentally in how the feed motion is achieved: in drilling, the tool '
    'itself performs the axial feed into the stationary workpiece, whereas in milling, the workpiece (via the '
    'table) provides a multi-axial feed relative to the rotating cutter.'
)
add_para(
    'The survey also revealed important differences in workpiece clamping strategies dictated by the kinematic '
    'requirements and cutting forces of each machine type. On lathes, workpieces are clamped in chucks or between '
    'centers to allow rotation. On drilling machines, vices, T-clamps, and specialized jigs hold the workpiece '
    'stationary on the table. On milling machines, the higher cutting forces necessitate more robust clamping using '
    'heavy-duty vices, T-slot clamp sets, and specialized accessories such as dividing heads and rotary tables, '
    'which also expand the geometric capabilities of the machine.'
)
add_para(
    'Furthermore, the tooling differs significantly across the three machine types. Lathes use single-point '
    'cutting tools mounted on the tool post, drills use multi-point axial tools (twist drills, reamers, taps) '
    'held in the spindle chuck or taper, and milling machines employ a wide variety of multi-point rotary cutters '
    '(end mills, face mills, slab mills, form cutters), each mounted according to the machine\'s orientation. '
    'The specialized variants of each machine — such as turret lathes, radial drilling machines, and CNC machining '
    'centers — demonstrate how the basic kinematic principles are adapted and expanded to meet industrial demands '
    'for speed, precision, flexibility, and heavy-duty machining.'
)
add_para('')
add_para('')

# Save the document
doc.save(SRC)
print(f"\nDocument saved successfully to: {SRC}")
print(f"Total paragraphs after edit: {len(doc.paragraphs)}")
