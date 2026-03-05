"""
Convert Lab 1 HTML report to DOCX — fixed spacing version.
Forces tight paragraph spacing on EVERY element.
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from bs4 import BeautifulSoup, NavigableString
from PIL import Image
import os, shutil

REPORT_DIR = r"C:\Users\stani\.gemini\antigravity\scratch\Manufacturing\Reports"
HTML_FILE = os.path.join(REPORT_DIR, "Report L1 Survey of machine tools for selected cutting processes.html")
DOCX_FILE = os.path.join(REPORT_DIR, "Report L1 Survey of machine tools for selected cutting processes.docx")
ONEDRIVE_DOCX = r"C:\Users\stani\OneDrive\Pulpit\PŁ_dokumenty\Manufacturing\Reports\Report L1 Survey of machine tools for selected cutting processes.docx"

# Convert JPEG-in-PNG files
for fn in sorted(os.listdir(REPORT_DIR)):
    if fn.startswith('fig') and fn.endswith('.png'):
        fpath = os.path.join(REPORT_DIR, fn)
        img = Image.open(fpath)
        if img.format != 'PNG':
            img.save(fpath, 'PNG')
        img.close()

with open(HTML_FILE, 'r', encoding='utf-8') as f:
    soup = BeautifulSoup(f.read(), 'lxml')

doc = Document()

# ---- FORCE tight spacing on Normal style ----
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(12)
normal.paragraph_format.space_after = Pt(0)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.line_spacing = 1.0
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

# Fix List Bullet style too
lb = doc.styles['List Bullet']
lb.font.name = 'Times New Roman'
lb.font.size = Pt(12)
lb.paragraph_format.space_after = Pt(0)
lb.paragraph_format.space_before = Pt(0)
lb.paragraph_format.line_spacing = 1.0

# Heading styles — black, tight spacing
for i in range(1, 4):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = 'Times New Roman'
    hs.font.color.rgb = RGBColor(0, 0, 0)
    hs.font.size = Pt([0, 16, 14, 13][i])
    hs.paragraph_format.space_before = Pt(12 if i == 1 else 8)
    hs.paragraph_format.space_after = Pt(4)
    hs.paragraph_format.line_spacing = 1.0

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)


def tighten(para, space_after=Pt(4), space_before=Pt(0)):
    """Force tight spacing on a paragraph."""
    para.paragraph_format.space_after = space_after
    para.paragraph_format.space_before = space_before
    para.paragraph_format.line_spacing = 1.15


def add_run(para, text, bold=False, italic=False, size=None, subscript=False, superscript=False):
    run = para.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = size or Pt(12)
    run.bold = bold
    run.italic = italic
    if subscript:
        run.font.subscript = True
    if superscript:
        run.font.superscript = True
    return run


def process_inline(para, el):
    if isinstance(el, NavigableString):
        text = str(el)
        if text.strip() or text == ' ':
            add_run(para, text)
    elif el.name in ('strong', 'b'):
        add_run(para, el.get_text(), bold=True)
    elif el.name in ('em', 'i'):
        add_run(para, el.get_text(), italic=True)
    elif el.name == 'sub':
        add_run(para, el.get_text(), subscript=True)
    elif el.name == 'sup':
        add_run(para, el.get_text(), superscript=True)
    else:
        # Recursively handle nested elements
        for child in el.children:
            process_inline(para, child)


def add_image(src):
    img_path = os.path.join(REPORT_DIR, src)
    if not os.path.exists(img_path):
        print(f"  MISSING: {src}")
        return
    try:
        doc.add_picture(img_path, width=Inches(4.5))
        p = doc.paragraphs[-1]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        tighten(p, space_after=Pt(2), space_before=Pt(6))
        print(f"  OK: {src}")
    except Exception as e:
        print(f"  ERROR: {src}: {e}")


# ====================== PARSE HTML ======================
body = soup.find('body')

for el in body.children:
    if isinstance(el, NavigableString):
        continue
    tag = el.name
    if not tag:
        continue

    # --- Title block ---
    if el.get('class') and 'title-block' in el.get('class', []):
        for child in el.children:
            if isinstance(child, NavigableString) or child.name == 'br':
                continue
            if child.name == 'p':
                text = child.get_text().strip()
                if not text:
                    continue
                cls = child.get('class', [])
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                tighten(p, space_after=Pt(1), space_before=Pt(1))
                if 'report-title' in cls:
                    add_run(p, text, bold=True, size=Pt(16))
                elif 'lab-title' in cls:
                    add_run(p, text, bold=True, size=Pt(14))
                elif 'course' in cls:
                    add_run(p, text, size=Pt(11))
                else:
                    for c in child.children:
                        process_inline(p, c)
        continue

    # --- Headings ---
    if tag in ('h1', 'h2', 'h3'):
        level = int(tag[1])
        doc.add_heading(el.get_text(), level=level)
        continue

    # --- Paragraphs ---
    if tag == 'p':
        text = el.get_text().strip()
        if not text:
            continue
        cls = el.get('class', [])
        p = doc.add_paragraph()
        if 'figure-caption' in cls:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, el.get_text(), italic=True, size=Pt(10))
            tighten(p, space_after=Pt(2), space_before=Pt(0))
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for child in el.children:
                process_inline(p, child)
            tighten(p, space_after=Pt(4))
        continue

    # --- Lists ---
    if tag == 'ul':
        for li in el.find_all('li', recursive=False):
            p = doc.add_paragraph(style='List Bullet')
            for child in li.children:
                process_inline(p, child)
            tighten(p, space_after=Pt(2))
        continue

    # --- Divs (figures / legends) ---
    if tag == 'div':
        cls = el.get('class', [])
        if 'figure-container' in cls:
            img = el.find('img')
            if img:
                add_image(img.get('src', ''))
            cap = el.find('p', class_='figure-caption')
            if cap:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_run(p, cap.get_text(), italic=True, size=Pt(10))
                tighten(p, space_after=Pt(2))
        elif 'legend' in cls:
            for ptag in el.find_all('p'):
                txt = ptag.get_text().strip()
                if not txt:
                    continue
                p = doc.add_paragraph(txt)
                p.paragraph_format.left_indent = Cm(1.5)
                tighten(p, space_after=Pt(0), space_before=Pt(0))
                for r in p.runs:
                    r.font.size = Pt(11)
                    r.font.name = 'Times New Roman'
        continue

# Save — use temp file to avoid lock if Word has it open
TEMP_DOCX = DOCX_FILE.replace('.docx', '_NEW.docx')
doc.save(TEMP_DOCX)
print(f"\nSaved to: {TEMP_DOCX}")
try:
    shutil.copy2(TEMP_DOCX, DOCX_FILE)
    print(f"Replaced: {DOCX_FILE}")
except PermissionError:
    print(f"NOTE: Could not overwrite original (open in Word?). Use the _NEW file.")
try:
    shutil.copy2(TEMP_DOCX, ONEDRIVE_DOCX)
    print(f"Copied to OneDrive: {ONEDRIVE_DOCX}")
except PermissionError:
    ONEDRIVE_NEW = ONEDRIVE_DOCX.replace('.docx', '_NEW.docx')
    shutil.copy2(TEMP_DOCX, ONEDRIVE_NEW)
    print(f"Saved new to OneDrive: {ONEDRIVE_NEW}")
print(f"Paragraphs: {len(doc.paragraphs)}")
print("DONE")
