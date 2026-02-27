"""
Script to add Quantitative Data and Team Collaboration sections to the ceramic mug production report.
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Open the document
doc_path = r"C:\Users\stani\.gemini\antigravity\scratch\Lab9_Excel\Production management - ceramic mug.docx"
doc = Document(doc_path)

# Add a page break before new sections
doc.add_page_break()

# =============================================
# SECTION 8: QUANTITATIVE DATA
# =============================================

heading8 = doc.add_heading('8. Quantitative Data', level=1)

intro_p = doc.add_paragraph()
intro_p.add_run('This section summarizes the key quantitative parameters and targets defined throughout the production process for ceramic mugs.')

# 8.1 Product Specifications
doc.add_heading('8.1 Product Specifications', level=2)
doc.add_paragraph('The following table presents the dimensional specifications for standard ceramic mugs produced in the facility:')

# Create specifications table
specs_table = doc.add_table(rows=9, cols=2)
specs_table.style = 'Table Grid'
specs_data = [
    ('Parameter', 'Value'),
    ('Capacity', '300–350 ml'),
    ('Height', '90–100 mm'),
    ('Outer Diameter', '80–85 mm'),
    ('Wall Thickness', '4–6 mm'),
    ('Handle Width', '30–35 mm'),
    ('Handle Thickness', '8–12 mm'),
    ('Total Weight', '350–450 g'),
    ('Base Diameter', '60–65 mm'),
]
for i, (param, value) in enumerate(specs_data):
    row = specs_table.rows[i]
    row.cells[0].text = param
    row.cells[1].text = value
    if i == 0:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

doc.add_paragraph()  # spacing

# 8.2 Process Parameters
doc.add_heading('8.2 Process Parameters', level=2)
doc.add_paragraph('Critical temperature and time parameters for each production stage:')

process_table = doc.add_table(rows=8, cols=2)
process_table.style = 'Table Grid'
process_data = [
    ('Process Parameter', 'Value'),
    ('Bisque Firing Temperature', '900–1060°C'),
    ('Glaze Firing Temperature', '>1200°C'),
    ('Bisque Firing Duration', '8–12 hours (incl. heating/cooling)'),
    ('Glaze Firing Duration', '10–14 hours (incl. heating/cooling)'),
    ('Controlled Drying Time', '24–48 hours'),
    ('Clay Moisture Content at Forming', '20–25%'),
    ('Bone-Dry Moisture Content', '<1%'),
]
for i, (param, value) in enumerate(process_data):
    row = process_table.rows[i]
    row.cells[0].text = param
    row.cells[1].text = value
    if i == 0:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

doc.add_paragraph()

# 8.3 Key Performance Indicators
doc.add_heading('8.3 Key Performance Indicators (KPIs)', level=2)
doc.add_paragraph('Target values for production efficiency and quality control:')

kpi_table = doc.add_table(rows=9, cols=3)
kpi_table.style = 'Table Grid'
kpi_data = [
    ('KPI Metric', 'Target', 'Description'),
    ('First Pass Yield (FPY)', '≥92%', 'Mugs passing final inspection without rework'),
    ('Cycle Time (Shaping)', '45 sec', 'Time to form one mug body'),
    ('Kiln Utilization', '≥95%', 'Volume used per firing cycle'),
    ('Defect Rate (Post-Firing)', '<3%', 'Cracks, glaze defects, or deformations'),
    ('On-Time Delivery Rate', '≥98%', 'Orders shipped within committed window'),
    ('Machine Downtime', '<2%', 'Unplanned downtime vs scheduled time'),
    ('Daily Production Capacity', '2000–3000', 'Mugs per 8-hour shift'),
    ('Kiln Batch Size', '200–500', 'Mugs per firing cycle'),
]
for i, row_data in enumerate(kpi_data):
    row = kpi_table.rows[i]
    for j, text in enumerate(row_data):
        row.cells[j].text = text
        if i == 0:
            for paragraph in row.cells[j].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

doc.add_paragraph()

# 8.4 Material Consumption
doc.add_heading('8.4 Material Consumption', level=2)
doc.add_paragraph('Estimated material requirements per mug:')

material_table = doc.add_table(rows=4, cols=2)
material_table.style = 'Table Grid'
material_data = [
    ('Material', 'Consumption per Mug'),
    ('Clay Body', '400–500 g (before 10–15% shrinkage)'),
    ('Glaze Application', '30–50 ml'),
    ('Energy (Firing)', '0.8–1.2 kWh'),
]
for i, (mat, cons) in enumerate(material_data):
    row = material_table.rows[i]
    row.cells[0].text = mat
    row.cells[1].text = cons
    if i == 0:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True

doc.add_paragraph()

# =============================================
# SECTION 9: TEAM COLLABORATION
# =============================================

doc.add_heading('9. Team Collaboration', level=1)

doc.add_paragraph('This report was developed collaboratively by a diverse, international team of five students. The work was divided according to the 5Ps framework to ensure comprehensive coverage of all aspects of ceramic mug production management.')

# 9.1 Task Distribution
doc.add_heading('9.1 Task Distribution', level=2)

doc.add_paragraph('The team organized the work through a shared communication channel where sections were assigned based on individual strengths and interests. The project was initiated on 21 January 2026 when the team established the section structure and agreed on responsibilities. The division of responsibilities was as follows:')

task_table = doc.add_table(rows=6, cols=3)
task_table.style = 'Table Grid'
task_data = [
    ('Team Member', 'Student ID', 'Assigned Sections'),
    ('Maciej Gara', '257093', 'Introduction, Product, People, Summary'),
    ('Stanisław Kotowicz', '257100', 'Process'),
    ('Michał Ogłuszka', '257104', 'Plant'),
    ('Sofia Neves', '905422', 'Programme (Scheduling, Resource Allocation)'),
    ('Ana Carvalho', '905431', 'Programme (Performance Monitoring, KPIs)'),
]
for i, row_data in enumerate(task_data):
    row = task_table.rows[i]
    for j, text in enumerate(row_data):
        row.cells[j].text = text
        if i == 0:
            for paragraph in row.cells[j].paragraphs:
                for run in paragraph.runs:
                    run.bold = True

doc.add_paragraph()

# 9.2 Collaborative Approach
doc.add_heading('9.2 Collaborative Approach', level=2)

doc.add_paragraph('Communication was maintained through a group messaging platform, allowing for real-time coordination and feedback exchange. Each member contributed their expertise to their assigned sections:')

p1 = doc.add_paragraph()
p1.add_run('Introduction and Product sections ').bold = True
p1.add_run('provided the foundational context for the production analysis.')

p2 = doc.add_paragraph()
p2.add_run('Process section ').bold = True
p2.add_run('detailed the step-by-step manufacturing workflow from raw materials to final packaging.')

p3 = doc.add_paragraph()
p3.add_run('Plant section ').bold = True
p3.add_run('covered facility design, equipment requirements, factory layout, and safety measures.')

p4 = doc.add_paragraph()
p4.add_run('Programme section ').bold = True
p4.add_run('addressed production scheduling strategies, resource allocation using pull systems and Kanban techniques, and performance monitoring through KPIs.')

p5 = doc.add_paragraph()
p5.add_run('People and Summary sections ').bold = True
p5.add_run('concluded the report by outlining workforce requirements and synthesizing the key findings.')

doc.add_paragraph()

# 9.3 Integration and Quality Assurance
doc.add_heading('9.3 Integration and Quality Assurance', level=2)

doc.add_paragraph('To ensure consistency across all sections, the team conducted reviews of each other\'s contributions and aligned formatting, terminology, and style. The final document represents a unified effort that demonstrates the practical application of the 5Ps framework to real-world production management.')

# Save the updated document
output_path = r"C:\Users\stani\.gemini\antigravity\scratch\Lab9_Excel\Production management - ceramic mug - FINAL.docx"
doc.save(output_path)

print(f"Document saved successfully to: {output_path}")
print("Added sections:")
print("  - 8. Quantitative Data (with 4 subsections and tables)")
print("  - 9. Team Collaboration (with 3 subsections)")
