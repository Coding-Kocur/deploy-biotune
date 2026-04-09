from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re

original_text = """Dzień 1 (push): 
• 3 min cardio na orbitreku/air bike do 140 HR max + rozgrzewka (z naciskiem na stożki rotatorów!) 
• Incline press (skos dodatni z hantlami) 3 serie, 8-12 repsów, RIR: 1, Tempo: 2/1/1/0 
• Dipy z obciążeniem 3s, 8-12 r, RIR: 1, 1/1/1/0 
• Katana extension (na warkoczach z wyciągu dolnego), 3s, 8-12 r, RIR: 0, 2/1/1/0 
• Rozpiętki (na bramie z góry lub na maszynie pec fly) 2s + dropset(~60% ciężaru na głównej, potem 40%), 8-15 r, RIR: 0, 2/1/2/1 
• Laterale – wzniosy bokiem pojedynczo na kablach 2 x 2s, 8-12 r, RIR: 0, 3/0/1/0 
 
Dzień 2 (pull): 
• 3 min cardio na orbitreku/air bike do 140 HR max + rozgrzewka 
• Lat pulldown na wąskim uchwycie 2s, 8-12 r, RIR: 1, 2/0/X/0 
• T-bar row, 4s, 5-8r (ciężko), RIR: 1, 2/1/1/0 
• Lat pullover (narciarz) 2s+dropset (~60% ciężaru na głównej, potem 40%), 8-12 r, RIR: 0, 2/1/1/0 
• Rear delt fly lub face pulls 3s, 10-15 r, RIR: 0, 2/0/X/0 
• Biceps curl z hantlami na incline (opierasz się plecami nie brzuchem), 3s, 8-12 r, RIR: 0, 3/0/1/0 
• Wzniosy nóg lub ab crunche na maszynie (brzuch) 3s, 10-15 r, RIR: 0, 2/1/1/0 
 
Dzień 3 (Nogi A): 
• 3 min rowerek / bieżnia pod górkę + rozgrzewka 
• Hack squat / suwnica (stopy jak najniżej żeby maksymalnie zaangażować czwórki ale bez odrywania pięt) 4s, 5-8 r, RIR: 2, 1/1/1/1 
• Leg extension – wyprosty na maszynie 3s+dropset (~60% ciężaru na głównej, potem 40%), 10-15 r, RIR: 0, 3/0/X/0 
• Bułgary z hantlami w oparciu o ławke, 2 x 2s, 10-15 r, RIR: 1, 2/1/1/0 
• Hip adductor (przywodzenie), 3 s, 10-15 r, RIR: 0, 3/0/1/1 
 
Dzień 4 – rest (jak masz siłe - 20-30min zone 2 cardio, pomaga w regeneracji po nogach) 
 
Dzień 5 (Upper): 
• 3 min cardio na orbitreku/air bike do 140 HR max + rozgrzewka (z naciskiem na stożki rotatorów!) 
• Podciąganie na drążku nadchwytem (plus obciążenie), 3s, 8-12 r, RIR: 1, 2/1/X/0 
• Wide press (na maszynie) lub wyciskanie na płaskiej z hantlami, 3s, 8-12 r, RIR: 1, 2/0/X/0 
• Wiosło szeroko na smith maszynie podchwytem lub cable rowy (też szeroko), 2s, 5-8 r (ciężko), RIR: 1, 2/1/1/0 
• OHP z hantlami na ławce, 2s, 6-10 r, RIR: 1, 2/0/1/0 
• Cross cable triceps extension, 2s, 8-12 r, RIR:0, 2/0/X/1 
• Biceps curl na EZ barze, 2s, 8-12 r, RIR: 0, 1/0/X/0 
 
Dzień 6 (Nogi B): 
• RDL ze sztangą, 3s, 6-10, RIR: 1, 3/1/1/0 
• Hip thrusty (najlepiej na maszynie), 3s, 8-12r, RIR: 2, 2/0/X/2 (przytrzymanie na górze) 
• Leg curl – uginanie na dwójki, 4s, 10-15 r, RIR: 1 (pierwsze dwie, potem do upadku), 2/0/1/0 
• Ławka rzymska, 3s, 10-15r, RIR: 2, 3/0/X/0 
• Calf raises na suwnicy 3s, 15-30 r, RIR: 2, 1/0/1/0 
 
Dzień 7 - Rest"""

subs_w1 = {
    'Incline press (skos dodatni z hantlami)': 'Incline press ze sztangą 45 stopni',
    'Katana extension (na warkoczach z wyciągu dolnego)': 'Triceps pushdown (wyciąg górny)',
    'Lat pulldown na wąskim uchwycie': 'Single arm cable row z wyciągu dolnego',
    'OHP z hantlami na ławce': 'Military press na smith',
    'Biceps curl na EZ barze': 'Biceps curl na modlitewniku',
    'Hip thrusty (najlepiej na maszynie)': 'Przysiad dzień dobry',
    'Hack squat / suwnica': 'Hack squat / pendulum squat',
    '(przytrzymanie na górze)': ''
}

subs_w2 = {
    'Hack squat / suwnica': 'Suwnica'
}

def write_bolded_bullet(doc, text):
    # Remove the literal bullet character from the text
    text = text.replace('• ', '').replace('•', '').strip()
    
    # Check if this is a cardio/warmup line (no bold required)
    if 'cardio' in text.lower() or 'rowerek' in text.lower():
        p = doc.add_paragraph(text, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25)
        return
        
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Inches(0.25)
    
    # Find where the exercise name ends and the sets/reps start.
    # Common splits: " 3s", " 4s", " 2s", " 3 serie", " 2 x 2s", " 3 s,"
    # We can use regex to find the first occurrence of a number followed by 's' or 'serie'
    match = re.search(r'(\s\d+\s*s[,+ ]|\s\d+\s+serie|\s\d+\s*x\s*\d+s)', text)
    
    if match:
        split_idx = match.start()
        exercise_name = text[:split_idx]
        rest_of_text = text[split_idx:]
        
        # Sometime there's a comma right before the split, let's include it in the bold or remove it
        if exercise_name.endswith(','):
             exercise_name = exercise_name[:-1]
             rest_of_text = ',' + rest_of_text
             
        p.add_run(exercise_name).bold = True
        p.add_run(rest_of_text)
    else:
        # If we can't cleanly split it by sets, just write it entirely
        p.add_run(text)

def create_training_plan():
    doc = Document()
    
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    sec = doc.sections[0]
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    
    title = doc.add_heading('Plan Treningowy - Maciej', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()
    
    lines = original_text.split('\n')
    
    # ------------------ TYDZIEŃ 1 ------------------
    doc.add_heading('Tydzień 1', level=1)
    doc.add_paragraph()

    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
            
        mod_line = line
        for old, new in subs_w1.items():
            if old in mod_line:
                mod_line = mod_line.replace(old, new)
        
        # Clean up double spacing resulting from empty string substitutions
        mod_line = mod_line.replace(' ,', ',').replace('  ', ' ')
                
        if mod_line.startswith('Dzień'):
            p = doc.add_paragraph()
            p.add_run(mod_line).bold = True
        elif mod_line.startswith('•'):
            write_bolded_bullet(doc, mod_line)
        else:
            doc.add_paragraph(mod_line)
            
    doc.add_page_break()

    # ------------------ TYDZIEŃ 2 ------------------
    doc.add_heading('Tydzień 2', level=1)
    doc.add_paragraph()

    for line in lines:
        line = line.strip()
        if not line:
            doc.add_paragraph()
            continue
            
        mod_line = line
        for old, new in subs_w2.items():
            if old in mod_line:
                mod_line = mod_line.replace(old, new)
                
        if mod_line.startswith('Dzień'):
            p = doc.add_paragraph()
            p.add_run(mod_line).bold = True
        elif mod_line.startswith('•'):
            write_bolded_bullet(doc, mod_line)
        else:
            doc.add_paragraph(mod_line)
            
    save_path = r'C:\Users\stani\OneDrive\Pulpit\Plany\Plan_Treningowy_Maciej.docx'
    doc.save(save_path)
    print(f'Successfully created file: {save_path}')

create_training_plan()
